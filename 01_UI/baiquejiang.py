# !/usr/bin/python
# -*- coding:utf-8 -*-

import os
import re
import sys
from docx import Document
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from GongZhongHao import GzhHandler
from Ui_MainWindow import Ui_MainWindow
from Drawer import Drawer

from enum import Enum
class ErrorType(Enum):
    NoError = 0
    NoFile = 1
    NoKeyWord = 2
    NoCoverImg = 3
    NoTemplate = 4
    NoAccessToken = 5
    NoCheckFormat = 6
    FileFormatError = 7

class Article():
    def __init__(self):
        self.cate = ""
        self.keyword = ""
        self.Profile = ""
        self.template = ""
        self.titles = []
        self.authors = []
        self.reviews = []
        self.contents = []
        self.comments = []
        self.addresses = [] 

def DeleteNumber(text):
    remove_chars = '[0-9’!"#$%&\'()*+,-./:;<=>?@，。?★、…【】《》？“”‘’！[\\]^_`{|}~]+'
    return re.sub(remove_chars, '', text)    

class BaiqueApp(QMainWindow, Ui_MainWindow):
    def __init__(self):
        super().__init__()
        AppId = "wx38330ee81eefe3da"
        AppSecret = "07b9b0ec8cab6746e208295616b80165"
        self.File = ""
        self.Article = ""
        self.Template = ""   
        self.setupUi(self)
        self.ResetClassMember()
        self.gzh = GzhHandler(AppId, AppSecret)
        self.ConnectWidgetToFun()
        self.ShowTemplatePng()   
        self.radioButton.setChecked(True)
    
    def EnableKeywordInput(self):
        if self.checkBox.isChecked():
            self.lineEdit_3.setEnabled(True)
        else:
            self.lineEdit_3.setEnabled(False)
    
    def ResetClassMember(self):
        self.Profile = ""      # 个人简介
        self.AwardList = []    # 奖次列表
        self.TitleList = []    # 题目列表
        self.AuthorList = []   # 作者列表
        self.CommentList = []  # 注释列表
        self.ContentList = []  # 内容列表
        self.AddressList = []  # 地址列表
        self.ReviewsList = []  # 评语列表
        self.ReviewerList = [] # 评委列表      
        self.textEdit_4.setText("评语的默认正则表达式为：..+评:.*|..+评：.*|【.+评】.*")
    
    # 接受GzhHandler的返回结果，准备标题、摘要，调用GzhHandler.uploadArticle
    @pyqtSlot(str)
    def PreviewFinish(self, Article):
        self.textEdit_4.setText("预览已生成，请点击preview.html查看！")
        self.Article = Article
        Year = str(self.spinBox_2.value())
        Month = str(self.spinBox_3.value())
        Quarter = self.comboBox_9.currentText()
        Category = "词"

        # 合成标题和摘要
        if self.Template == "个人专辑":
            title = "白雀奖 || XXX诗词曲作品专辑"
            digest = self.Profile[:100] # 摘要限制字数
        elif self.Template == "月度入围":
            title = f"白雀奖 || 20{Year}年第{Quarter}季度{Month}月份{Category}部入围作品及点评"
            digest = "本期评委："
            for Reviewer in self.ReviewerList:
                digest += Reviewer.replace("评","")+"，"
                digest += "本期共有"+str(len(self.TitleList))+"首作品入围。"
        elif self.Template == "季度获奖":
            title = "白雀奖 || 20"+Year+"年第"+Quarter+"季度"+Category+"部获奖公告"
            digest = "终评评委："
            for Reviewer in self.ReviewerList:
                digest += Reviewer.replace("评","")+"，"
        
        self.gzh.uploadArticle(title,"cover.jpg",digest,self.Article)
    
    # 接受GzhHandler的返回结果，在窗口打印信息
    @pyqtSlot(str)
    def UploadFinish(self, Result):
        self.textEdit_4.setText(Result)
    
    def ShowTemplatePng(self):
        png_1 = QPixmap("个人专辑.png")
        self.label_16.setPixmap(png_1)
        png_2 = QPixmap("月度入围.png")
        self.label_17.setPixmap(png_2)
        png_3 = QPixmap("季度获奖.png")
        self.label_18.setPixmap(png_3)
    
    def ConnectWidgetToFun(self):
        self.pushButton.clicked.connect(self.SelectFile)
        self.pushButton_2.clicked.connect(self.GenOneCertificate)
        self.pushButton_3.clicked.connect(self.GenAllCertificate)
        self.pushButton_4.clicked.connect(self.CheckDocxFormat_1)
        self.pushButton_5.clicked.connect(self.SelectFile)
        self.pushButton_6.clicked.connect(self.CheckDocxFormat_2)
        self.pushButton_7.clicked.connect(self.GenPreview)
        self.pushButton_8.clicked.connect(self.ClearOutputWindow)
        self.pushButton_9.clicked.connect(self.ClearOutputWindow)
        self.gzh.signal_1.connect(self.PreviewFinish)
        self.gzh.signal_2.connect(self.UploadFinish)
        self.radioButton.toggled.connect(lambda:self.SetComboBoxList())
        self.radioButton_2.toggled.connect(lambda:self.SetComboBoxList())
        self.radioButton_3.toggled.connect(lambda:self.SetComboBoxList())
    
    def StackedWidgetSwitch(self):
        item = self.treeWidget.currentItem()
        if item.text(0) == "一键生成奖状":
            self.stackedWidget.setCurrentIndex(0)
        elif item.text(0) == "一键公众号文章":
            self.stackedWidget.setCurrentIndex(1)
    
    def SelectFile(self):
        fname = QFileDialog.getOpenFileName(None, "选择文档", "/", "*.docx")
        if fname[0]:
            self.File = fname[0]
            self.textEdit_3.setText("已选中："+self.File)
            self.textEdit_4.setText("已选中："+self.File)
    
    def GenAllCertificate(self):
        Year = self.spinBox.value()
        Stamp = self.comboBox_7.currentText()
        Quarter = self.comboBox_6.currentText()
        Category = self.comboBox_2.currentText()
        if self.File == "":
            self.textEdit_3.setText("请先导入文档！")
        else:
            self.ParseDocx3()
            if len(self.TitleList) == len(self.AuthorList) == len(self.ContentList) == len(self.AwardList):
                self.textEdit_3.setText("文档格式无误，正在自动生成所有奖状！")
                myDrawer = Drawer(Stamp,Category,Quarter,Year)
                for i in range(len(self.TitleList)):
                    result = myDrawer.DrawCertificate(self.AwardList[i], self.TitleList[i], self.AuthorList[i], self.ContentList[i])
                    self.textEdit_3.setText("正在生成：" + result)
            else:
                self.textEdit_3.setText("文档格式不正确，请选择需要打印的输出信息！")
    
    def GenOneCertificate(self):
        Year = self.spinBox.value()
        Author = self.lineEdit.text()
        Title = self.lineEdit_2.text()
        Award = self.comboBox.currentText()
        Content = self.textEdit.toPlainText()
        Stamp = self.comboBox_7.currentText()
        Quarter = self.comboBox_6.currentText()
        Category = self.comboBox_2.currentText()
        if Content == "" or Author == "" or Title == "":
            self.textEdit_3.setText("请输入必要的作品信息！")
        else:
            myDrawer = Drawer(Stamp,Category,Quarter,Year)
            myDrawer.DrawCertificate(Award,Title,Author,Content)  
    
    def GenPreview(self):
        result = ErrorType.NoError
        Keyword = self.lineEdit_3.text()
        Category = self.comboBox_8.currentText()
        article = Article()
        article.cate = Category
        article.keyword = ""
        article.profile = self.Profile
        article.template = self.Template
        article.titles = self.TitleList
        article.authors = self.AuthorList
        article.reviews = self.ReviewsList
        article.contents = self.ContentList
        article.comments = self.CommentList
        article.addresses = self.AddressList
        
        if result == ErrorType.NoError and self.File == "":
            result = ErrorType.NoFile
            self.textEdit_4.setText("请先导入文档！")
        if result == ErrorType.NoError and self.checkBox.isChecked() and Keyword == "":
            result = ErrorType.NoKeyWord
            self.textEdit_4.setText("请输入搜索图片关键词！")
        if result == ErrorType.NoError and self.gzh.getAccessToken() != "True":
            result = ErrorType.NoAccessToken
            self.textEdit_4.setText(self.gzh.getAccessToken())
        if result == ErrorType.NoError and len(self.TitleList) == 0:
            result = ErrorType.NoCheckFormat
            self.textEdit_4.setText("请先点击'检查格式'按钮！")
        if result == ErrorType.NoError:
            if self.Template == "个人专辑":
                if len(self.TitleList) != len(self.ContentList) != len(self.CommentList):
                    result = ErrorType.FileFormatError
                    self.textEdit_4.setText("文档格式不正确，请选择需要打印的输出信息！")
            elif self.Template == "月度入围":
                if len(self.TitleList) != len(self.ReviewsList) != len(self.ContentList) != len(self.CommentList):
                    result = ErrorType.FileFormatError
                    self.textEdit_4.setText("文档格式不正确，请选择需要打印的输出信息！")
            elif self.Template == "季度获奖":
                if len(self.TitleList) != len(self.ReviewsList) != len(self.ContentList) != len(self.CommentList) != len(self.AddressList) != len(self.AuthorList):
                    result = ErrorType.FileFormatError
                    self.textEdit_4.setText("文档格式不正确，请选择需要打印的输出信息！")
        if result == ErrorType.NoError and os.path.isfile("cover.jpg") == False:
            result = ErrorType.NoCoverImg
            self.textEdit_4.setText("请添加封面图片！")
        if result == ErrorType.NoError:
            self.textEdit_4.setText("正在生成文章并上传公众号...")
            self.gzh.generateHtml(article)
    
    # 检查奖状文章格式
    def CheckDocxFormat_1(self):
        if self.File == "":
            self.textEdit_3.setText("请先导入文档！")
        else:
            self.ParseDocx3()
        
            PrintText = ""
            Class = self.comboBox_5.currentText()
            if Class == "标题":
                for item in self.TitleList:
                    PrintText += item + "\n" 
                PrintText += "数量为：" + str(len(self.TitleList))
            elif Class == "内容":
                for item in self.ContentList:
                    PrintText += item + "\n\n"
                PrintText += "数量为：" + str(len(self.ContentList))
            elif Class == "作者":
                for item in self.AuthorList:
                    PrintText += item + "\n\n"
                PrintText += "数量为：" + str(len(self.AuthorList))
            
            self.textEdit_3.setText(PrintText)
    
    # 检查公众号文章格式
    def CheckDocxFormat_2(self):
        if self.File == "":
            self.textEdit_4.setText("请先导入文档！")
        else:
            if self.radioButton.isChecked() == True:
                self.Template = "个人专辑"
                self.ParseDocx1()
            elif self.radioButton_2.isChecked() == True:
                self.Template = "月度入围"
                self.ParseDocx2()
            elif self.radioButton_3.isChecked() == True:
                self.Template = "季度获奖"
                self.ParseDocx3()
        
            PrintText = ""
            Class = self.comboBox_4.currentText()
            if Class == "标题":
                for item in self.TitleList:
                    PrintText += item + "\n" 
                PrintText += "数量为：" + str(len(self.TitleList))
            if Class == "内容":
                for item in self.ContentList:
                    PrintText += item + "\n\n"
                PrintText += "数量为：" + str(len(self.ContentList))
            if Class == "注释":
                for item in self.CommentList:
                    PrintText += item + "\n"
                PrintText += "数量为：" + str(len(self.CommentList))
            if Class == "评语":
                for item in self.ReviewsList:
                    PrintText += item + "\n\n"
                PrintText += "数量为：" + str(len(self.ReviewsList))
            if Class == "作者":
                for item in self.AuthorList:
                    PrintText += item + "\n"
                PrintText += "数量为：" + str(len(self.AuthorList))
            if Class == "地址":
                for item in self.AddressList:
                    PrintText += item + "\n"
                PrintText += "数量为：" + str(len(self.AddressList))
            if Class == "简介":
                PrintText = self.Profile
            self.textEdit_4.setText(PrintText)
    
    def SetComboBoxList(self):
        _translate = QCoreApplication.translate
        self.comboBox_4.clear()
        self.comboBox_4.addItem("")
        self.comboBox_4.addItem("")
        self.comboBox_4.setItemText(0, _translate("MainWindow", "标题"))
        self.comboBox_4.setItemText(1, _translate("MainWindow", "内容"))
        if self.radioButton_2.isChecked() == True:
            self.comboBox_4.addItem("")
            self.comboBox_4.addItem("")
            self.comboBox_4.setItemText(2, _translate("MainWindow", "注释"))
            self.comboBox_4.setItemText(3, _translate("MainWindow", "评语"))
        elif self.radioButton_3.isChecked() == True:
            self.comboBox_4.addItem("")
            self.comboBox_4.addItem("")
            self.comboBox_4.addItem("")
            self.comboBox_4.addItem("")
            self.comboBox_4.setItemText(2, _translate("MainWindow", "注释"))
            self.comboBox_4.setItemText(3, _translate("MainWindow", "评语"))
            self.comboBox_4.setItemText(4, _translate("MainWindow", "作者"))
            self.comboBox_4.setItemText(5, _translate("MainWindow", "地址"))
        elif self.radioButton.isChecked() == True:
            self.comboBox_4.addItem("")
            self.comboBox_4.setItemText(2, _translate("MainWindow", "简介"))
    
    def ClearOutputWindow(self):
        self.textEdit_3.setText("")
        self.textEdit_4.setText("")
    
    # 解析个人专辑
    def ParseDocx1(self):
        
        # 定义正则表达式
        RegProfile = "个人简介"
        RegComment = "^注.*"
        RegContent = "正文"
        
        # 定义局部变量
        Comment = "" # 单条注释
        Content = "" # 单条内容
        TitleFlag = False # 题目以下是作品正文
        
        # 开始解析
        self.ResetClassMember()
        document = Document(self.File)
        for paragraph in document.paragraphs:
            if paragraph.text != "":
                FindProfile = re.search(RegProfile, paragraph.text)
                FindContent = re.search(RegContent, paragraph.text)
                FindComment = re.search(RegComment, paragraph.text)
                if FindProfile:
                    self.Profile = paragraph.text
                else:
                    if TitleFlag == True:
                        self.TitleList.append(paragraph.text)
                        TitleFlag = False
                    elif FindComment:
                        Comment += paragraph.text + "-"
                    else:
                        Content += paragraph.text + "-"
            else:
                TitleFlag = True
                if Content != "":
                    self.ContentList.append(Content)                   
                    self.CommentList.append(Comment)                   
                    Comment = ""
                    Content = ""
    
    # 解析月度入围名单
    def ParseDocx2(self):
        
        # 定义正则表达式
        RegComment = "^注.*"
        RegReviews = "..+评:.*|..+评：.*|【.+评】.*"
        RegTitle = "[0-9]+、.*|[0-9]+-[1-3]+、.*|[0-9]+·.*|[0-9]+-[1-3]+·.*"
        
        # 定义局部变量
        Reviews = "" # 单条评语
        Content = "" # 单条内容
        Comment = "" # 单条注释
        TitleFlag = False  # 题目以下是作品正文
        SingleReview = [2] # 用于将评语拆分成两部分：[评委, 评语]
        
        # 开始解析
        self.ResetClassMember()
        document = Document(self.File)
        for paragraph in document.paragraphs:
            if paragraph.text != "":
                FindTitle = re.search(RegTitle, paragraph.text)
                FindReviews = re.search(RegReviews, paragraph.text)
                FindComment = re.search(RegComment, paragraph.text)
                if FindReviews:
                    TitleFlag = False
                    if Content != "":
                        self.ContentList.append(Content)
                        Content = ""
                    # 单条评语拆分成列表：[评委, 评语]
                    SingleReview = re.split(":|：", FindReviews.group())
                    if SingleReview[0] not in self.ReviewerList:
                        self.ReviewerList.append(SingleReview[0])
                    # 给评委名字前后加上装饰符：【】
                    if "【" not in SingleReview[0]:
                        Reviews += "【" + SingleReview[0] + "】" + SingleReview[1] + "-"
                    else:
                        Reviews += "-" + FindReviews.group()
                elif FindTitle:
                    self.TitleList.append(FindTitle.group())
                    TitleFlag = True
                elif TitleFlag == True and FindComment == None:
                    Content += paragraph.text + "-"
                elif FindComment:
                    Comment += paragraph.text + "-"
            else:
                if Reviews != "":
                    self.CommentList.append(Comment)
                    self.ReviewsList.append(Reviews)
                    Comment = ""
                    Reviews = ""
    
    # 解析获奖名单
    def ParseDocx3(self):
        
        # 定义正则表达式
        RegComment = "^注.*"
        RegReviews = "..+评:.*|..+评：.*|【.+评】.*"
        RegAuthor = "作者：.*|作者:.*|姓名：.*|姓名:.*"
        RegAddress = "地址:.*|地址：.*|住址:.*|住址：.*"
        RegTitle = "[0-9]+、.*|[0-9]+-[1-3]+、.*|[0-9]+·.*|[0-9]+-[1-3]+·.*"
        
        # 定义局部变量
        Reviews = "" # 单条评语
        Content = "" # 单条内容
        Comment = "" # 单条注释
        AddressFlag = False # 地址以下是作品正文
        CurrentAward = ""   # 记录当前获奖级别
        SingleReview = [2]  # 用于将评语拆分成两部分：[评委, 评语]
        AwardClass = ["一等奖", "二等奖", "三等奖", "优秀奖"] # 获奖级别
        idx = 0 # 用于遍历获奖级别列表
        
        # 开始解析
        self.ResetClassMember()
        document = Document(self.File)
        for paragraph in document.paragraphs:
            if paragraph.text != "":
                FindTitle = re.search(RegTitle, paragraph.text)
                FindAuthor = re.search(RegAuthor, paragraph.text)
                FindAddress = re.search(RegAddress, paragraph.text)
                FindReviews = re.search(RegReviews, paragraph.text)
                FindComment = re.search(RegComment, paragraph.text)
                FindAward = re.search(AwardClass[idx], paragraph.text)
                if FindReviews:
                    AddressFlag = False
                    if Content != "":
                        self.ContentList.append(Content)
                        Content = ""                    
                    # 将评语拆分成两部分：[评委, 评语]
                    SingleReview = re.split(":|：", FindReviews.group())
                    if SingleReview[0] not in self.ReviewerList:
                        self.ReviewerList.append(SingleReview[0])
                    # 给评委名字前后加上修饰符：【】
                    if "【" not in SingleReview[0]:
                        Reviews += "【" + SingleReview[0] + "】" + SingleReview[1] + "-"
                    else:
                        Reviews += "-" + FindReviews.group()
                elif FindAward:
                    CurrentAward = AwardClass[idx]
                    if idx + 1 < 4:
                        idx += 1
                elif FindTitle:
                    self.TitleList.append(DeleteNumber(FindTitle.group()))
                elif FindAuthor:
                    self.AuthorList.append(FindAuthor.group()[3:])
                elif FindAddress:
                    self.AddressList.append(FindAddress.group()[3:])
                    AddressFlag = True
                elif AddressFlag == True and FindComment == None:
                    Content += paragraph.text + "-"
                elif FindComment:
                    Comment += paragraph.text + "-"
            else:
                if Reviews != "":
                    self.AwardList.append(CurrentAward)                   
                    self.ReviewsList.append(Reviews)
                    self.CommentList.append(Comment)
                    Reviews = ""
                    Comment = ""
        # for item in AwardClass:
            # self.AwardList.append(tempAwardList.count(item))

if __name__ == '__main__':
  app = QApplication(sys.argv)
  baique = BaiqueApp()
  baique.show()
  sys.exit(app.exec_())